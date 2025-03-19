import tkinter as tk
from tkinter import filedialog, messagebox
from collections import Counter
import docx
import magic
import zipfile
import os
import tempfile

class KeywordAnalyzerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("关键词频率分析器")
        
        # 文件选择
        self.file_label = tk.Label(root, text="选择Word文档:")
        self.file_label.grid(row=0, column=0, padx=5, pady=5)
        
        self.file_entry = tk.Entry(root, width=40)
        self.file_entry.grid(row=0, column=1, padx=5, pady=5)
        
        self.browse_button = tk.Button(root, text="浏览", command=self.browse_file)
        self.browse_button.grid(row=0, column=2, padx=5, pady=5)
        
        # 关键词输入
        self.keywords_label = tk.Label(root, text="输入关键词（用逗号分隔）:")
        self.keywords_label.grid(row=1, column=0, padx=5, pady=5)
        
        self.keywords_entry = tk.Entry(root, width=40)
        self.keywords_entry.grid(row=1, column=1, padx=5, pady=5)
        
        # 分析按钮
        self.analyze_button = tk.Button(root, text="分析", command=self.analyze_document)
        self.analyze_button.grid(row=2, column=1, padx=5, pady=5)
        
        # 结果显示
        self.result_text = tk.Text(root, height=10, width=50)
        self.result_text.grid(row=3, column=0, columnspan=3, padx=5, pady=5)
        
    def browse_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word文档", "*.docx *.docm")])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, file_path)
        
    def analyze_document(self):
        file_path = self.file_entry.get()
        keywords = self.keywords_entry.get().split(',')
        
        if not file_path:
            messagebox.showerror("错误", "请选择一个Word文档")
            return
            
        if not keywords:
            messagebox.showerror("错误", "请输入至少一个关键词")
            return
            
        try:
            mime = magic.Magic(mime=True)
            file_type = mime.from_file(file_path)
            
            if file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                           'application/vnd.ms-word.document.macroEnabled.main+xml']:
                doc = docx.Document(file_path)
                text = " ".join([para.text for para in doc.paragraphs])
            elif file_type == 'application/zip':
                # 处理zip文件
                with tempfile.TemporaryDirectory() as tmpdir:
                    with zipfile.ZipFile(file_path, 'r') as zip_ref:
                        zip_ref.extractall(tmpdir)
                    
                    # 显示zip文件结构
                    print("\n压缩包内容：")
                    for root, dirs, files in os.walk(tmpdir):
                        level = root.replace(tmpdir, '').count(os.sep)
                        indent = ' ' * 4 * (level)
                        print(f'{indent}{os.path.basename(root)}/')
                        subindent = ' ' * 4 * (level + 1)
                        for f in files:
                            print(f'{subindent}{f}')
                    
                    # 递归查找zip中的所有Word文档
                    text = ""
                    for root, dirs, files in os.walk(tmpdir):
                        for file in files:
                            if file.endswith('.docx') or file.endswith('.docm'):
                                doc_path = os.path.join(root, file)
                                try:
                                    doc = docx.Document(doc_path)
                                    text += " ".join([para.text for para in doc.paragraphs]) + " "
                                    print(f"成功读取文件: {file}")
                                except Exception as e:
                                    print(f"警告：无法读取文件 {file}: {str(e)}")
                    
                    if not text.strip():
                        raise ValueError("压缩包中未找到可读取的Word文档")
            else:
                raise ValueError(f"不支持的文件类型: {file_type}")
            
            # 统计关键词频率
            word_count = Counter(text.lower().split())
            keyword_freq = {kw.strip().lower(): word_count.get(kw.strip().lower(), 0) for kw in keywords}
            
            # 显示结果
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, "关键词频率统计结果:\n\n")
            for kw, freq in keyword_freq.items():
                self.result_text.insert(tk.END, f"{kw}: {freq} 次\n")
                
        except Exception as e:
            messagebox.showerror("错误", f"无法读取文件: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = KeywordAnalyzerApp(root)
    root.mainloop()
